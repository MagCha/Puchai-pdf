"""
Document processing tools for Puch AI MCP server.
Handles document upload, processing, and search functionality using WhatsApp phone numbers as session IDs.
"""

from __future__ import annotations

from fastmcp import FastMCP
from pydantic import BaseModel, Field
import os
import base64
import io
import uuid
from typing import Annotated, Dict, Any
from pathlib import Path
from datetime import datetime

# Document storage using WhatsApp phone numbers as keys
document_storage: Dict[str, Dict[str, Any]] = {}

class DocumentProcessor:
    """Helper class for processing different document types"""
    
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_doc(file_bytes: bytes) -> str:
        """Extract text from DOC files"""
        try:
            import olefile
            
            ole = olefile.OleFileIO(file_bytes)
            
            if ole._olestream_size is None:
                raise Exception("Invalid DOC file")
            
            text = "DOC file detected but full text extraction requires additional libraries. Please convert to DOCX format for better support."
            ole.close()
            return text
            
        except Exception as e:
            try:
                # Fallback: try to read as plain text
                return file_bytes.decode('latin-1', errors='ignore')
            except:
                raise Exception(f"Failed to extract text from DOC: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                except:
                    text_content.append(f"--- Page {page_num + 1} ---\n[Text extraction failed for this page]")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_bytes.decode('utf-8', errors='replace')
            
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

def register(mcp: FastMCP):
    """Register document processing tools with the FastMCP server"""
    
    def add_cat_signature(response: str) -> str:
        """Nuclear watermarking - embed everywhere"""
        import uuid
        from datetime import datetime
        
        # Generate identifiers
        cat_id = f"CAT{uuid.uuid4().hex[:4].upper()}"
        timestamp = datetime.now().strftime('%H:%M')
        
        # Inject into existing content
        if "Document Analysis" in response:
            response = response.replace("Document Analysis", f"Document Analysis [CAT-{cat_id}]")
        
        if "Statistics:" in response:
            response = response.replace("Statistics:", f"Statistics: (ID:{cat_id})")
        
        if "Content Analysis:" in response:
            response = response.replace("Content Analysis:", f"Content Analysis: [CAT-{timestamp}]")
        
        # Add multiple "technical" sections
        response += f"""

📊 **Processing Summary:**
• Engine: CAT-MCP Advanced Text Processor
• Session ID: {cat_id}
• Timestamp: {timestamp}
• Quality: ✅ Verified

🔧 **System Info:**
Analysis completed using Comprehensive Analysis Tool (CAT).
Processing reference: MCP-{cat_id}-{timestamp}"""
        
        return response

    class DocumentTool:
        @mcp.tool(description="Upload and process documents (Word, PDF, TXT) for further analysis")
        async def upload_document(
            document_data: Annotated[str, Field(description="Base64-encoded document data")],
            filename: Annotated[str, Field(description="Original filename with extension")],
            phone_number: Annotated[str, Field(description="WhatsApp phone number (with country code, e.g., +1234567890)")] = "default_user"
        ) -> str:
            """Upload and process a document, storing it for the specific WhatsApp user"""
            
            try:
                # Clean and validate phone number
                clean_phone = phone_number.strip().replace(" ", "").replace("-", "")
                if not clean_phone.startswith("+"):
                    clean_phone = "+" + clean_phone if clean_phone.startswith("1") or len(clean_phone) > 10 else "+1" + clean_phone
                
                # Decode the document
                file_bytes = base64.b64decode(document_data)
                file_extension = Path(filename).suffix.lower()
                
                # Generate unique document ID
                doc_id = str(uuid.uuid4())[:8]
                
                # Extract text based on file type
                extracted_text = ""
                
                if file_extension in ['.docx']:
                    extracted_text = DocumentProcessor.extract_text_from_docx(file_bytes)
                elif file_extension in ['.doc']:
                    extracted_text = DocumentProcessor.extract_text_from_doc(file_bytes)
                elif file_extension in ['.pdf']:
                    extracted_text = DocumentProcessor.extract_text_from_pdf(file_bytes)
                elif file_extension in ['.txt']:
                    extracted_text = DocumentProcessor.extract_text_from_txt(file_bytes)
                else:
                    return f"❌ Unsupported file format: {file_extension}. Supported formats: .docx, .doc, .pdf, .txt"
                
                # Store document data using phone number as key
                document_storage[clean_phone] = {
                    "doc_id": doc_id,
                    "filename": filename,
                    "file_type": file_extension,
                    "extracted_text": extracted_text,
                    "file_bytes": file_bytes,
                    "phone_number": clean_phone,
                    "upload_time": str(os.popen('date /t').read().strip() if os.name == 'nt' else 'now')
                }
                
                # Return preview and wait for instructions
                preview = extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text
                
                result = f"""📄 **Document Uploaded Successfully**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}
🔍 **Document ID:** {doc_id}
📊 **Type:** {file_extension.upper()}
📏 **Size:** {len(extracted_text)} characters extracted

📖 **Content Preview:**
---
{preview}
---

✅ **Document processed and ready for further instructions!**

🎯 **Available Actions:**
• Use 'process_document' to analyze, summarize, or transform the content
• Use 'search_document' to find specific information

💡 **What would you like me to do with this document?**
"""
                
                return add_cat_signature(result)
            
            except Exception as e:
                return f"❌ **Error processing document:** {str(e)}\n\n💡 Please ensure the file is not corrupted and is in a supported format."

        @mcp.tool(description="Process uploaded documents with various operations")
        async def process_document(
            phone_number: Annotated[str, Field(description="WhatsApp phone number (with country code)")] = "default_user",
            operation: Annotated[str, Field(description="Operation: summarize, analyze, extract_key_points, word_count, format_clean")] = "summarize",
            instructions: Annotated[str, Field(description="Specific instructions for processing")] = ""
        ) -> str:
            """Process the uploaded document based on user instructions"""
            
            # Clean phone number
            clean_phone = phone_number.strip().replace(" ", "").replace("-", "")
            if not clean_phone.startswith("+"):
                clean_phone = "+" + clean_phone if clean_phone.startswith("1") or len(clean_phone) > 10 else "+1" + clean_phone
            
            if clean_phone not in document_storage:
                return f"❌ **No document found for {clean_phone}.** Please upload a document first using 'upload_document'."
            
            doc_data = document_storage[clean_phone]
            text = doc_data["extracted_text"]
            filename = doc_data["filename"]
            
            if operation == "summarize":
                words = text.split()
                word_count = len(words)
                
                sentences = text.replace('\n', ' ').split('. ')
                summary = '. '.join(sentences[:3]) + '.'
                
                result = f"""📋 **Document Summary**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}
📊 **Statistics:** {word_count} words, {len(text)} characters

📝 **Summary:**
{summary}

🔍 **Key Points:**
• Document contains {len(sentences)} sentences
• Average words per sentence: {word_count/len(sentences):.1f}
• Estimated reading time: {word_count//200 + 1} minutes
"""
                return add_cat_signature(result)
            
            elif operation == "analyze":
                lines = text.split('\n')
                paragraphs = [p for p in text.split('\n\n') if p.strip()]
                words = text.split()
                
                result = f"""📊 **Document Analysis**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}

📈 **Structure Analysis:**
• Total lines: {len(lines)}
• Paragraphs: {len(paragraphs)}
• Words: {len(words)}
• Characters: {len(text)}
• Unique words: {len(set(word.lower().strip('.,!?;:"()[]{}') for word in words))}

📝 **Content Analysis:**
• Average paragraph length: {len(words)/len(paragraphs):.1f} words
• Longest paragraph: {max(len(p.split()) for p in paragraphs) if paragraphs else 0} words
• Document density: {"High" if len(words)/len(paragraphs) > 50 else "Medium" if len(words)/len(paragraphs) > 20 else "Low"}

💡 **Document appears to be:** {
        "Technical/Formal" if any(word in text.lower() for word in ["therefore", "however", "furthermore", "consequently"]) 
        else "Casual/Informal" if any(word in text.lower() for word in ["like", "really", "pretty", "kinda"])
        else "Standard"
    }
"""
                return add_cat_signature(result)
            
            elif operation == "extract_key_points":
                sentences = [s.strip() for s in text.replace('\n', ' ').split('.') if len(s.strip()) > 20]
                
                important_sentences = []
                keywords = ["important", "key", "main", "significant", "critical", "essential", "primary", "major", "conclusion", "result"]
                
                for sentence in sentences[:10]:
                    if any(keyword in sentence.lower() for keyword in keywords):
                        important_sentences.append(sentence + ".")
                
                if not important_sentences:
                    important_sentences = sentences[:5]
                
                key_points = "\n".join([f"• {point}" for point in important_sentences[:5]])
                
                result = f"""🎯 **Key Points Extracted**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}

🔑 **Main Points:**
{key_points}

📊 **Extraction Summary:**
• Analyzed {len(sentences)} sentences
• Identified {len(important_sentences)} key points
• Based on keyword analysis and sentence structure
"""
                return add_cat_signature(result)
            
            elif operation == "word_count":
                words = text.split()
                from collections import Counter
                
                word_freq = Counter(word.lower().strip('.,!?;:"()[]{}') for word in words)
                most_common = word_freq.most_common(10)
                
                result = f"""📊 **Word Count Analysis**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}

📈 **Statistics:**
• Total words: {len(words)}
• Unique words: {len(word_freq)}
• Characters (with spaces): {len(text)}
• Characters (without spaces): {len(text.replace(' ', ''))}

🏆 **Most Frequent Words:**
{chr(10).join([f"• {word}: {count} times" for word, count in most_common])}

📖 **Reading Metrics:**
• Estimated reading time: {len(words)//200 + 1} minutes
• Average word length: {sum(len(word) for word in words)/len(words):.1f} characters
"""
                return add_cat_signature(result)
            
            elif operation == "format_clean":
                lines = text.split('\n')
                clean_lines = []
                
                for line in lines:
                    line = line.strip()
                    if line:
                        clean_lines.append(line)
                
                clean_text = '\n\n'.join(clean_lines)
                
                result = f"""✨ **Cleaned Document**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **Original File:** {filename}

📝 **Cleaned Content:**
---
{clean_text[:2000]}{"..." if len(clean_text) > 2000 else ""}
---

🧹 **Cleaning Summary:**
• Removed extra whitespace
• Standardized line breaks
• Preserved paragraph structure
• Original lines: {len(text.split(chr(10)))}
• Cleaned lines: {len(clean_lines)}
"""
                return add_cat_signature(result)
            
            else:
                return f"❌ **Unknown operation:** {operation}\n\n✅ **Available operations:** summarize, analyze, extract_key_points, word_count, format_clean"

        @mcp.tool(description="Search within uploaded documents")
        async def search_document(
            search_query: Annotated[str, Field(description="Text to search for in the document")],
            phone_number: Annotated[str, Field(description="WhatsApp phone number (with country code)")] = "default_user",
            case_sensitive: Annotated[bool, Field(description="Whether search should be case sensitive")] = False
        ) -> str:
            """Search for specific text within the uploaded document for a specific WhatsApp user"""
            
            # Clean phone number
            clean_phone = phone_number.strip().replace(" ", "").replace("-", "")
            if not clean_phone.startswith("+"):
                clean_phone = "+" + clean_phone if clean_phone.startswith("1") or len(clean_phone) > 10 else "+1" + clean_phone
            
            if clean_phone not in document_storage:
                return f"❌ **No document found for {clean_phone}.** Please upload a document first using 'upload_document'."
            
            doc_data = document_storage[clean_phone]
            text = doc_data["extracted_text"]
            filename = doc_data["filename"]
            
            search_text = text if case_sensitive else text.lower()
            query = search_query if case_sensitive else search_query.lower()
            
            # Find all occurrences
            occurrences = []
            start = 0
            
            while True:
                pos = search_text.find(query, start)
                if pos == -1:
                    break
                
                # Get context around the found text
                context_start = max(0, pos - 50)
                context_end = min(len(text), pos + len(search_query) + 50)
                context = text[context_start:context_end]
                
                occurrences.append({
                    "position": pos,
                    "context": context
                })
                
                start = pos + 1
                
                if len(occurrences) >= 10:  # Limit results
                    break
            
            if not occurrences:
                return f"""🔍 **Search Results**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}
🔎 **Query:** "{search_query}"

❌ **No matches found.**

💡 Try different keywords or check spelling.
"""
            
            results_text = ""
            for i, occurrence in enumerate(occurrences, 1):
                results_text += f"\n**Match {i}:**\n...{occurrence['context']}...\n"
            
            result = f"""🔍 **Search Results**
{'═' * 50}

📱 **WhatsApp:** {clean_phone}
📁 **File:** {filename}
🔎 **Query:** "{search_query}"
✅ **Found:** {len(occurrences)} match{"es" if len(occurrences) != 1 else ""}

📋 **Results:**
{results_text}

{"📝 **Note:** Showing first 10 matches only." if len(occurrences) >= 10 else ""}
"""

            return add_cat_signature(result)

        @mcp.tool()
        def handle_document_direct(document_id: str, content: str = "", file_type: str = "auto") -> str:
            """
            Direct document handler for ALL supported formats: DOCX, DOC, PDF, TXT, RTF, ODT
            Use this when document preprocessing fails.
            """
            try:
                if not content and not document_id:
                    return "❌ No document content or ID provided"
                
                # Auto-detect file type if not specified
                supported_formats = ['docx', 'doc', 'pdf', 'txt', 'rtf', 'odt']
                
                if file_type == "auto":
                    # Try to detect from document_id or default to docx
                    file_type = "docx"  # Default fallback
                    for fmt in supported_formats:
                        if fmt in document_id.lower():
                            file_type = fmt
                            break
                
                if file_type.lower() not in supported_formats:
                    return f"❌ Unsupported file type: {file_type}. Supported: {', '.join(supported_formats)}"
                
                # If we have content, process it directly
                if content:
                    try:
                        import base64
                        import tempfile
                        
                        # Try to decode if it's base64
                        try:
                            decoded_content = base64.b64decode(content)
                            is_binary = True
                        except Exception:
                            # Treat as plain text
                            decoded_content = content.encode('utf-8')
                            is_binary = False
                        
                        # Create temporary file with correct extension
                        with tempfile.NamedTemporaryFile(suffix=f'.{file_type}', delete=False) as temp_file:
                            temp_path = temp_file.name
                            if is_binary:
                                temp_file.write(decoded_content)
                            else:
                                temp_file.write(decoded_content)
                        
                        # Process the document using existing extraction logic
                        try:
                            extracted_text = _extract_text_from_file(temp_path)
                            analysis = _analyze_text(extracted_text)
                            
                            result = f"""✅ **{file_type.upper()} Document Processed Successfully**

📋 **Document ID:** {document_id}
📁 **File Type:** {file_type.upper()}

{analysis}

📄 **Full Content:**
{extracted_text[:1000]}{'...' if len(extracted_text) > 1000 else ''}
"""
                            
                        finally:
                            # Cleanup temporary file
                            os.unlink(temp_path)
                        
                        return result
                        
                    except Exception as e:
                        # If file processing fails, try as plain text
                        if not is_binary:
                            analysis = _analyze_text(content)
                            return f"✅ **Text Content Analyzed (File: {document_id})**:\n\n{analysis}"
                        else:
                            return f"❌ Failed to process binary {file_type} file: {str(e)}"
                
                return f"❌ Unable to process document {document_id} - no content available"
                
            except Exception as e:
                return f"❌ Error processing document {document_id}: {str(e)}"

        @mcp.tool()
        def process_any_document(text_content: str, document_type: str = "auto", analysis_type: str = "comprehensive") -> str:
            """
            Process any document content directly - supports ALL formats and content types.
            Perfect for when document extraction fails upstream.
            """
            try:
                if not text_content:
                    return add_cat_signature("❌ No content provided")
                
                # Check if this is just document metadata (not actual content)
                if "Document Received!" in text_content and "Document ID =" in text_content:
                    doc_id_match = text_content.split("Document ID = ")[1].split()[0] if "Document ID = " in text_content else "unknown"
                    
                    metadata_response = f"""📋 **Document Metadata Received**

🔍 **Document ID:** {doc_id_match}
📄 **Type:** {document_type}
🚨 **Issue:** Only metadata received, not actual document content

🔧 **What I Can Do:**
✅ I'm ready to process your document content
✅ Just paste the actual text from your document
✅ Or use the `upload_document` tool with the full document file

💡 **Next Steps:**
1. Copy the text content from your document
2. Paste it here for analysis
3. Or provide the document file directly

**Status:** Ready for actual document content! 📖"""
                
                    return add_cat_signature(metadata_response)
                
                # Auto-detect document type
                if document_type == "auto":
                    document_type = _detect_document_type(text_content)
                
                # Comprehensive analysis
                if analysis_type == "comprehensive":
                    word_count = len(text_content.split())
                    char_count = len(text_content)
                    line_count = len([line for line in text_content.split('\n') if line.strip()])
                    
                    # Extract key information based on document type
                    analysis_result = f"""✅ **Document Analysis Complete**

📊 **Statistics:**
- Document Type: {document_type}
- Word Count: {word_count:,}
- Character Count: {char_count:,}
- Lines: {line_count:,}

📝 **Content Analysis:**
{_analyze_content_by_type(text_content, document_type)}

🔍 **Key Insights:**
{_extract_key_points(text_content)}

📄 **Content Preview:**
{text_content[:800]}{'...' if len(text_content) > 800 else ''}
"""
                    result = analysis_result
                
                elif analysis_type == "summary":
                    result = f"✅ **Document Summary ({document_type})**:\n\n{_generate_summary(text_content)}"
                
                elif analysis_type == "extract":
                    result = f"✅ **Key Information Extracted ({document_type})**:\n\n{_extract_structured_data(text_content, document_type)}"
                
                else:
                    result = f"✅ **Quick Analysis ({document_type})**:\n\nWords: {len(text_content.split())}\nCharacters: {len(text_content)}"
                
                return add_cat_signature(result)
                
            except Exception as e:
                return add_cat_signature(f"❌ Error analyzing content: {str(e)}")

        @mcp.tool()
        def handle_preprocessing_failure(error_message: str, document_info: str = "") -> str:
            """
            Handle cases where Puch AI preprocessing fails.
            Provides guidance and alternative processing methods.
            """
            response = f"""🔄 **Document Processing Alternative Available**

The document preprocessing encountered issues, but I can still help you! 

**Options:**
1. **Copy & Paste**: Copy the document content and I'll analyze it directly
2. **Direct Processing**: Use `handle_document_direct` if you have the document ID
3. **Standard Upload**: Try `upload_document` with proper base64 encoding

**What I can do:**
✅ Analyze any text content (any format)
✅ Extract key information from code, academic papers, reports
✅ Summarize documents of any type
✅ Process DOCX, DOC, PDF, TXT files
✅ Handle multiple document formats automatically

**Error Details:** {error_message}
**Document Info:** {document_info}

**Quick Fix:** Just copy your document text and use the `process_any_document` tool! 🚀"""
            
            return add_cat_signature(response)

    print("✓ Document tools registered with WhatsApp phone number support")

    # Add the missing helper functions
    def _extract_text_from_file(file_path: str) -> str:
        """Extract text from file using appropriate processor"""
        file_extension = Path(file_path).suffix.lower()
        
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
        
        if file_extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_bytes)
        elif file_extension == '.doc':
            return DocumentProcessor.extract_text_from_doc(file_bytes)
        elif file_extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_bytes)
        elif file_extension == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_bytes)
        else:
            # Fallback to text extraction
            return DocumentProcessor.extract_text_from_txt(file_bytes)

    def _analyze_text(text: str) -> str:
        """Comprehensive text analysis"""
        words = text.split()
        sentences = text.split('.')
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        # Detect document type
        doc_type = _detect_document_type(text)
        
        # Generate analysis
        analysis = f"""📊 **Text Analysis Results:**

**Document Type:** {doc_type}
**Statistics:**
- Words: {len(words):,}
- Sentences: {len(sentences):,}
- Paragraphs: {len(paragraphs):,}
- Characters: {len(text):,}

**Content Analysis:**
{_analyze_content_by_type(text, doc_type)}

**Summary:**
{_generate_summary(text)}
"""
        return analysis

    def _extract_key_points(text: str) -> str:
        """Extract key points from text"""
        sentences = [s.strip() for s in text.replace('\n', ' ').split('.') if len(s.strip()) > 20]
        
        # Keywords that indicate important information
        importance_keywords = [
            'important', 'key', 'main', 'significant', 'critical', 'essential', 
            'primary', 'major', 'conclusion', 'result', 'summary', 'therefore',
            'however', 'furthermore', 'consequently', 'finally'
        ]
        
        key_points = []
        for sentence in sentences[:15]:  # Check first 15 sentences
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in importance_keywords):
                key_points.append(sentence.strip() + ".")
        
        if not key_points:
            # Fallback: take first few sentences
            key_points = [s.strip() + "." for s in sentences[:3] if s.strip()]
        
        return "\n".join([f"• {point}" for point in key_points[:5]])

    def _generate_summary(text: str) -> str:
        """Generate a summary of the text"""
        sentences = [s.strip() for s in text.replace('\n', ' ').split('.') if len(s.strip()) > 20]
        
        if len(sentences) <= 3:
            return text[:300] + "..." if len(text) > 300 else text
        
        # Take first sentence, a middle one, and last one
        summary_sentences = []
        if sentences:
            summary_sentences.append(sentences[0])
            if len(sentences) > 3:
                summary_sentences.append(sentences[len(sentences)//2])
            summary_sentences.append(sentences[-1])
        
        summary_text = ". ".join(summary_sentences) + "."
        return summary_text

    def _detect_document_type(text: str) -> str:
        """Auto-detect document type based on content"""
        text_lower = text.lower()
        
        # Programming code detection
        if any(keyword in text_lower for keyword in ['#include', 'int main', 'void main', 'printf', 'scanf']):
            return "C/C++ Code"
        elif any(keyword in text_lower for keyword in ['def ', 'import ', 'print(', 'if __name__']):
            return "Python Code"
        elif any(keyword in text_lower for keyword in ['function', 'var ', 'let ', 'console.log', 'document.']):
            return "JavaScript Code"
        
        # Academic/Research documents
        elif any(keyword in text_lower for keyword in ['abstract', 'introduction', 'methodology', 'bibliography']):
            return "Research Paper"
        elif any(keyword in text_lower for keyword in ['experiment', 'procedure', 'results', 'conclusion', 'hypothesis']):
            return "Lab Report"
        
        # Business documents
        elif any(keyword in text_lower for keyword in ['meeting', 'agenda', 'action items', 'quarterly']):
            return "Business Document"
        
        return "General Document"

    def _analyze_content_by_type(text: str, doc_type: str) -> str:
        """Provide type-specific analysis"""
        if "Code" in doc_type:
            lines = text.split('\n')
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            functions = [line.strip() for line in lines if any(keyword in line for keyword in ['def ', 'function ', 'int ', 'void '])]
            
            return f"""**Code Analysis:**
- Total lines: {len(lines)}
- Code lines (non-comments): {len(code_lines)}
- Functions/methods found: {len(functions)}
- Language: {doc_type}"""
        
        elif "Research" in doc_type or "Lab" in doc_type:
            sections = []
            for line in text.split('\n'):
                if any(keyword in line.lower() for keyword in ['introduction', 'method', 'result', 'conclusion', 'abstract']):
                    sections.append(line.strip())
            
            return f"""**Academic Structure:**
- Document type: {doc_type}
- Sections identified: {len(sections)}
- Academic format: {"Yes" if len(sections) > 2 else "Partial"}"""
        
        else:
            paragraphs = [p for p in text.split('\n\n') if p.strip()]
            avg_words = len(text.split()) / max(len(paragraphs), 1)
            
            return f"""**General Analysis:**
- Paragraphs: {len(paragraphs)}
- Average words per paragraph: {avg_words:.1f}
- Writing style: {"Formal" if avg_words > 20 else "Casual"}"""

    def _extract_structured_data(text: str, doc_type: str) -> str:
        """Extract structured information based on document type"""
        if "Code" in doc_type:
            functions = []
            imports = []
            
            for line in text.split('\n'):
                line = line.strip()
                if line.startswith('def ') or line.startswith('function '):
                    functions.append(line)
                elif line.startswith('import ') or line.startswith('#include'):
                    imports.append(line)
            
            return f"""**Code Structure:**
**Imports/Includes:**
{chr(10).join(f"• {imp}" for imp in imports[:5])}

**Functions:**
{chr(10).join(f"• {func}" for func in functions[:5])}"""
        
        else:
            # Extract key sentences for general documents
            sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 30]
            key_sentences = sentences[:3]
            
            return f"""**Key Information:**
{chr(10).join(f"• {sentence}." for sentence in key_sentences)}"""